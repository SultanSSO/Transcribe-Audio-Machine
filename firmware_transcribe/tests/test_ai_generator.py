"""Test AI notulen generator — semua panggilan LLM di-mock (tanpa router asli).

Mencakup: load transkrip (TXT/JSON), build prompt (termasuk pemotongan),
generate_notulen_docx end-to-end (DOCX + TXT preview + callbacks), dan
MdToDocx.from_text (render DOCX dari string, tanpa file .md).
"""

import json
import os

import pytest

from src.notulen import ai_generator
from src.notulen.ai_generator import (
    MAX_INPUT_CHARS,
    build_prompt,
    generate_notulen_docx,
    load_transcript,
)


# ── fixtures ────────────────────────────────────────────────────────────
@pytest.fixture
def hasil_folder(tmp_path):
    """Folder hasil berisi transkrip.txt + transkrip.json + file lama .md."""
    folder = tmp_path / "01_rapat-contoh"
    folder.mkdir()
    (folder / "transkrip.txt").write_text(
        "Kita sepakat sistem diluncurkan bulan depan. Tim BTD akan koordinasi "
        "dengan PTT untuk data. Keputusan: pakai OJS versi terbaru.",
        encoding="utf-8",
    )
    (folder / "transkrip.json").write_text(json.dumps({
        "segments": [{"start": 0.0, "end": 2.0, "text": "Kita sepakat."}],
        "audio_duration_s": 75.0,
    }), encoding="utf-8")
    # file MD lama — harus diabaikan oleh alur AI
    (folder / "Notulen_rapat-contoh.md").write_text("# lama", encoding="utf-8")
    return folder


@pytest.fixture
def fake_llm(monkeypatch):
    """call_llm di-mock: kembalikan notulen markdown palsu."""
    def _fake(system, user):
        assert "=== MULAI TRANSKRIP ===" in user
        assert "JANGAN menambah, menebak" in user
        return (
            "# Notulensi Rapat Contoh\n\n"
            "## Ringkasan Eksekutif\n\n"
            "- Rapat membahas peluncuran sistem.\n\n"
            "## Keputusan\n\n"
            "1. Pakai OJS versi terbaru.\n"
        )
    monkeypatch.setattr(ai_generator, "call_llm", _fake)
    return _fake


# ── load_transcript ─────────────────────────────────────────────────────
def test_load_transcript_txt(hasil_folder):
    text, meta = load_transcript(str(hasil_folder))
    assert "sistem diluncurkan" in text
    assert meta["source_name"] == "01_rapat-contoh"


def test_load_transcript_fallback_json(tmp_path):
    folder = tmp_path / "02_tanpa-txt"
    folder.mkdir()
    (folder / "transkrip.json").write_text(json.dumps({
        "segments": [{"start": 0.0, "end": 1.0, "text": "Satu"},
                     {"start": 1.0, "end": 2.0, "text": "dua"}],
        "audio_duration_s": 2.0,
    }), encoding="utf-8")
    text, meta = load_transcript(str(folder))
    assert text == "Satu dua"
    assert meta["audio_duration"] == 2.0


def test_load_transcript_missing(tmp_path):
    with pytest.raises(FileNotFoundError):
        load_transcript(str(tmp_path))


# ── build_prompt ────────────────────────────────────────────────────────
def test_build_prompt(hasil_folder):
    text, meta = load_transcript(str(hasil_folder))
    system, user = build_prompt(text, meta)
    assert "notulis rapat profesional" in system
    assert "=== MULAI TRANSKRIP ===" in user
    assert "01_rapat-contoh" in user


def test_build_prompt_truncates_long_transcript():
    long_text = "kalimat panjang " * 100_000  # > MAX_INPUT_CHARS
    _, user = build_prompt(long_text, {"source_name": "01_rapat-besar",
                                       "audio_duration": None})
    assert "CATATAN: transkrip asli lebih panjang" in user
    # teks terkirim terpotong ke batas
    start = user.index("=== MULAI TRANSKRIP ===") + len("=== MULAI TRANSKRIP ===")
    end = user.index("=== AKHIR TRANSKRIP ===")
    assert len(user[start:end].strip()) <= MAX_INPUT_CHARS


# ── generate_notulen_docx (end-to-end, LLM di-mock) ────────────────────
def test_generate_notulen_docx(hasil_folder, fake_llm):
    events = {"logs": [], "progress": [], "finished": None}

    def on_log(m): events["logs"].append(m)
    def on_progress(p, m): events["progress"].append(p)
    def on_finished(r): events["finished"] = r

    result = generate_notulen_docx(str(hasil_folder), {
        "on_log": on_log, "on_progress": on_progress, "on_finished": on_finished,
    })

    assert result["docx_path"].endswith(".docx")
    assert os.path.isfile(result["docx_path"])
    assert os.path.isfile(result["txt_path"])
    assert result["model"] == ai_generator.AI_MODEL

    # DOCX valid — bisa dibuka python-docx dan berisi teks notulen
    from docx import Document
    doc = Document(result["docx_path"])
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Notulensi Rapat Contoh" in full
    assert "Ringkasan Eksekutif" in full

    # TXT preview: markdown di-strip, header ada, MD lama tidak tersentuh
    plain = open(result["txt_path"], encoding="utf-8").read()
    assert "NOTULEN AI" in plain
    assert "#" not in plain  # heading markdown sudah di-strip

    # callbacks
    assert events["finished"] == result
    assert any("DOCX" in m for m in events["logs"])
    assert events["progress"] and events["progress"][-1] == 100

    # MD lama tetap ada tapi tidak di-overwrite (fitur MD mati, bukan dihapus)
    assert (hasil_folder / "Notulen_rapat-contoh.md").is_file()
    # dan tidak ada file .md baru dibuat
    assert not any(p.suffix == ".md" for p in hasil_folder.iterdir()
                   if p.name != "Notulen_rapat-contoh.md")


def test_generate_notulen_docx_error_no_transcript(tmp_path, fake_llm):
    with pytest.raises(FileNotFoundError):
        generate_notulen_docx(str(tmp_path))


# ── MdToDocx.from_text ──────────────────────────────────────────────────
def test_mdtodocx_from_text(tmp_path):
    from src.export.docx_converter import MdToDocx
    out = tmp_path / "hasil.docx"
    md = (
        "# Judul Rapat\n\n"
        "## Bagian Satu\n\n"
        "- poin pertama\n"
        "- poin kedua\n\n"
        "## Bagian Dua\n\n"
        "1. aksi satu\n"
        "2. aksi dua\n"
    )
    MdToDocx.from_text(md, str(out))
    assert out.is_file()

    from docx import Document
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Judul Rapat" in texts
    assert any("poin pertama" in t for t in texts)


# ── call_llm error handling ─────────────────────────────────────────────
def test_call_llm_http_error(monkeypatch):
    class FakeResponse:
        status_code = 500
        text = "internal error"

    class FakeClient:
        def __init__(self, *a, **k): pass
        def __enter__(self): return self
        def __exit__(self, *a): return False
        def post(self, *a, **k): return FakeResponse()

    monkeypatch.setattr(ai_generator.httpx, "Client", FakeClient)
    with pytest.raises(ConnectionError) as ei:
        ai_generator.call_llm("sys", "usr")
    assert "9router" in str(ei.value) or "AI" in str(ei.value)


def test_call_llm_empty_response(monkeypatch):
    class FakeResponse:
        status_code = 200
        text = ""
        def json(self):
            return {"choices": [{"message": {"content": "   "}}]}

    class FakeClient:
        def __init__(self, *a, **k): pass
        def __enter__(self): return self
        def __exit__(self, *a): return False
        def post(self, *a, **k): return FakeResponse()

    monkeypatch.setattr(ai_generator.httpx, "Client", FakeClient)
    with pytest.raises(ValueError, match="Respons LLM kosong"):
        ai_generator.call_llm("sys", "usr")
