"""
Notulen MD → DOCX Converter (Generic)
======================================
Reads a Markdown notulen file and converts it to a professionally formatted
Word (.docx) document with styled headings, tables, bullets, and cover page.

Usage:
    # Activate venv first
    .venv\\Scripts\\Activate.ps1

    # Convert a specific MD file
    python generate_docx.py --md transcribe_hasil/03_rapat-koordinasi-event/Notulen_Rapat_Koordinasi_Event.md

    # Auto-find Notulen_*.md inside a result folder
    python generate_docx.py --folder 03_rapat-koordinasi-event

    # Convert all folders that have Notulen_*.md
    python generate_docx.py --all

    # Specify custom output path
    python generate_docx.py --md file.md --output hasil.docx
"""

import argparse
import os
import re
import subprocess
import sys
from pathlib import Path

# ═══════════════════════════════════════════════════════════════════════════
# Auto-bootstrap: re-run this script with the project venv if the current
# python doesn't have python-docx installed (e.g. system python won).
# ═══════════════════════════════════════════════════════════════════════════
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.abspath(os.path.join(_SCRIPT_DIR, "..", ".."))
_VENV_PY = os.path.join(_PROJECT_ROOT, ".venv", "Scripts", "python.exe")


def _rerun_with_venv_python():
    """Re-execute this script using .venv/Scripts/python.exe."""
    venv_prefix = os.path.abspath(os.path.join(_PROJECT_ROOT, ".venv"))
    if os.path.abspath(sys.prefix) == venv_prefix:
        return  # already running under the project venv

    if not os.path.isfile(_VENV_PY):
        print(
            f"[bootstrap] python saat ini bukan venv proyek ({sys.executable})\n"
            f"[bootstrap] dan venv tidak ditemukan di {_VENV_PY}.\n"
            f"[bootstrap] Aktifkan venv dulu: .\\.venv\\Scripts\\Activate.ps1",
            file=sys.stderr,
        )
        return  # let the original error surface

    print(
        f"[bootstrap] python saat ini bukan venv proyek ({sys.executable})\n"
        f"[bootstrap] menjalankan ulang dengan venv: {_VENV_PY}",
        file=sys.stderr,
    )
    # subprocess (bukan os.execv) agar path dengan spasi di-quote benar di Windows
    sys.exit(subprocess.call([_VENV_PY] + sys.argv))


_rerun_with_venv_python()

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ═══════════════════════════════════════════════════════════════════════════════
#  STYLING CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════

FONT_NAME = "Calibri"
FONT_SIZE = Pt(10)
HEADER_BG = "1B3A5C"       # navy
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
INFO_CELL_BG = "E8EDF2"    # light grey-blue

COLOR_H1 = RGBColor(0x1B, 0x3A, 0x5C)
COLOR_H2 = RGBColor(0x2E, 0x5C, 0x8A)
COLOR_H3 = RGBColor(0x3A, 0x6E, 0xA5)


# ═══════════════════════════════════════════════════════════════════════════════
#  LOW-LEVEL DOCX HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def shade_cell(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_run(paragraph, text, bold=False, italic=False, size=None, color=None):
    """Add a run with formatting to a paragraph. Returns the run."""
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def new_paragraph(doc, alignment=None, spacing_after=4, spacing_before=0):
    """Create a new paragraph with defaults."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.space_before = Pt(spacing_before)
    if alignment is not None:
        p.alignment = alignment
    return p


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a table with navy header row and styled data rows."""
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        make_run(p, str(text), bold=True, color=HEADER_FG)
        shade_cell(cell, HEADER_BG)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            make_run(p, str(text))

    # Normalize cell fonts
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(9) if len(rows) > 8 else Pt(10)

    return table


# ═══════════════════════════════════════════════════════════════════════════════
#  INLINE MARKDOWN PARSER  (bold, italic, code)
# ═══════════════════════════════════════════════════════════════════════════════

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"\*(.+?)\*")
BOLD_ITALIC_RE = re.compile(r"\*\*\*(.+?)\*\*\*")
CODE_RE = re.compile(r"`(.+?)`")
STRIP_RE = re.compile(r"^[\*\s]+|[\*\s]+$")  # leftover asterisks in plain text


def parse_inline(paragraph, text):
    """Parse inline markdown formatting and add runs to paragraph."""
    # Order matters: bold-italic first, then bold, then italic
    patterns = [
        (BOLD_ITALIC_RE, True, True),
        (BOLD_RE, True, False),
        (ITALIC_RE, False, True),
    ]

    # Simple tokenizer: split by patterns and interleave styled runs
    tokens = [(text, False, False)]  # (text, bold, italic)

    for pattern, bold, italic in patterns:
        new_tokens = []
        for t, b, i in tokens:
            if b or i:
                new_tokens.append((t, b, i))
            else:
                parts = pattern.split(t)
                for idx, part in enumerate(parts):
                    if idx % 2 == 0:
                        new_tokens.append((part, False, False))
                    else:
                        new_tokens.append((part, bold, italic))
        tokens = new_tokens

    for t, b, i in tokens:
        if not t:
            continue
        # Remove backtick code markers (simple approach)
        clean = CODE_RE.sub(r"\1", t)
        make_run(paragraph, clean, bold=b, italic=i)


# ═══════════════════════════════════════════════════════════════════════════════
#  MARKDOWN → DOCX CONVERTER
# ═══════════════════════════════════════════════════════════════════════════════

class MdToDocx:
    """Converts a Markdown file to a formatted .docx document."""

    def __init__(self, md_path=None, output_path=None, text=None):
        if text is not None:
            # Mode string: tidak ada file .md — DOCX langsung dari teks.
            self.md_path = None
            self._text = text
            if not output_path:
                raise ValueError("output_path wajib diisi saat memakai mode text")
        else:
            self.md_path = Path(md_path)
            if not self.md_path.exists():
                raise FileNotFoundError(f"Markdown file not found: {md_path}")
            self._text = None

        if output_path:
            self.output_path = Path(output_path)
        else:
            self.output_path = self.md_path.with_suffix(".docx")

        self.doc = Document()
        self._setup_page()
        self.lines = []
        self.idx = 0
        self.title = ""
        self.subtitle = ""

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        style = self.doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE

    # ── reading helpers ──────────────────────────────────────────────────

    def _peek(self, offset=0):
        i = self.idx + offset
        return self.lines[i] if i < len(self.lines) else ""

    def _advance(self):
        self.idx += 1

    def _consume(self):
        line = self.lines[self.idx]
        self.idx += 1
        return line

    def _is_table_sep(self, line):
        """Check if line is a markdown table separator: |---|----|"""
        return bool(re.match(r"^\|[\s\-:|]+\|$", line))

    def _is_table_row(self, line):
        return line.startswith("|") and line.rstrip().endswith("|")

    def _is_hr(self, line):
        return re.match(r"^[-*_]{3,}\s*$", line)

    def _is_bullet(self, line):
        return re.match(r"^[-*]\s", line)

    def _is_numbered(self, line):
        return re.match(r"^\d+\.\s", line)

    def _is_empty(self, line):
        return line.strip() == ""

    # ── table parsing ────────────────────────────────────────────────────

    def _parse_table(self):
        """Parse a markdown table starting at current index. Returns (headers, rows) or None."""
        start = self.idx
        rows_raw = []

        # Collect all consecutive table rows
        while self.idx < len(self.lines):
            line = self.lines[self.idx]
            if self._is_table_sep(line):
                self.idx += 1
                continue
            if self._is_table_row(line):
                cells = [c.strip() for c in line.strip("|").split("|")]
                rows_raw.append(cells)
                self.idx += 1
            else:
                break

        if not rows_raw:
            return None

        # First row = header if more than 1 row, else just data
        if len(rows_raw) >= 2 and len(rows_raw[0]) >= 2:
            # Check if the first row cells are ALL bold (likely header)
            first_row_has_bold = all("**" in c for c in rows_raw[0])
            if first_row_has_bold:
                headers = [re.sub(r"\*\*", "", c) for c in rows_raw[0]]
                data = rows_raw[1:]
                return headers, data

        # No clear header → treat first row as header
        headers = rows_raw[0]
        data = rows_raw[1:]
        if not data:
            # single-row table — treat as key-value
            return headers, []

        return headers, data

    def _is_info_table(self, headers, rows):
        """Detect if a 2-column table is a metadata/info table (no numeric header like #)."""
        if len(headers) != 2:
            return False
        if not rows:
            return True  # single-row 2-col = likely info
        # If first header is empty or starts with **, it's likely an info table
        h0 = headers[0].strip()
        if h0 == "" or h0 == "#":
            return False  # empty or "#" number — regular table
        return True

    def _render_info_table(self, headers, rows):
        """Render a 2-col key-value info table with shaded labels."""
        all_rows = [headers] + rows if rows else [headers]
        table = self.doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_data in all_rows:
            row = table.add_row()
            # First cell (label)
            cell0 = row.cells[0]
            cell0.text = ""
            p0 = cell0.paragraphs[0]
            parse_inline(p0, row_data[0])
            for run in p0.runs:
                run.bold = True
                run.font.size = Pt(10)
            shade_cell(cell0, INFO_CELL_BG)
            cell0.width = Cm(4)

            # Second cell (value)
            val = row_data[1] if len(row_data) > 1 else ""
            cell1 = row.cells[1]
            cell1.text = ""
            p1 = cell1.paragraphs[0]
            parse_inline(p1, val)

        # Font normalize
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)

    def _render_data_table(self, headers, rows):
        """Render a regular data table with styled header."""
        add_styled_table(self.doc, headers, rows)

    # ── heading rendering ────────────────────────────────────────────────

    def _render_h1(self, text):
        # Skip if it's the main title (already on cover)
        pass

    def _render_h2(self, text):
        h = self.doc.add_heading(text, level=2)
        for run in h.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(15)
            run.font.color.rgb = COLOR_H2

    def _render_h3(self, text):
        h = self.doc.add_heading(text, level=3)
        for run in h.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(13)
            run.font.color.rgb = COLOR_H3

    def _render_h4(self, text):
        h = self.doc.add_heading(text, level=4)
        for run in h.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_H3

    # ── main conversion ──────────────────────────────────────────────────

    def _build_cover(self):
        """Add cover page from H1/H2 and info table."""
        # Spacer
        for _ in range(4):
            self.doc.add_paragraph()

        # Title (H1)
        p = new_paragraph(self.doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        make_run(p, self.title, bold=True, size=Pt(22), color=COLOR_H1)

        # Subtitle (H2)
        if self.subtitle:
            p = new_paragraph(self.doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            make_run(p, self.subtitle, size=Pt(16), color=COLOR_H2)

        self.doc.add_paragraph()

    def _build_toc(self):
        """Build a manual table of contents from H2 headings."""
        self.doc.add_page_break()
        toc_title = self.doc.add_heading("Daftar Isi", level=1)
        for run in toc_title.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(15)
            run.font.color.rgb = COLOR_H1

        for line in self.lines:
            m = re.match(r"^##\s+(.+)$", line)
            if m:
                heading_text = m.group(1).strip()
                p = new_paragraph(self.doc, spacing_after=2)
                make_run(p, heading_text)

        self.doc.add_page_break()

    def convert(self):
        """Main conversion driver."""
        # Read all lines (mode text langsung dari string, tanpa file .md)
        if self._text is not None:
            self.lines = [line.rstrip() for line in self._text.splitlines()]
        else:
            with open(self.md_path, "r", encoding="utf-8") as f:
                self.lines = [line.rstrip() for line in f.readlines()]

        # ── PASS 1: Extract title & subtitle ──
        for line in self.lines:
            m1 = re.match(r"^#\s+(.+)", line)
            m2 = re.match(r"^##\s+(.+)", line)
            if m1 and not self.title:
                self.title = m1.group(1).strip()
            elif m2 and not self.subtitle:
                candidate = m2.group(1).strip()
                # Only treat very first H2 as subtitle if it comes right after H1
                if candidate and not re.match(r"^\d+\.", candidate) and "Kesimpulan" not in candidate.lower():
                    self.subtitle = candidate

        # ── PASS 2: Build cover before processing content ──
        self._build_cover()

        # ── PASS 3: Parse body line-by-line ──
        self.idx = 0
        first_section = True
        info_table_rendered = False

        while self.idx < len(self.lines):
            line = self.lines[self.idx]

            # ── skip empty ──
            if self._is_empty(line):
                self._advance()
                continue

            # ── horizontal rule → page break ──
            if self._is_hr(line):
                if first_section:
                    # First HR after cover info = build TOC then continue
                    self._build_toc()
                    first_section = False
                self._advance()
                continue

            # ── headings ──
            m_h1 = re.match(r"^#\s+(.+)", line)
            m_h2 = re.match(r"^##\s+(.+)", line)
            m_h3 = re.match(r"^###\s+(.+)", line)
            m_h4 = re.match(r"^####\s+(.+)", line)

            if m_h1:
                text = m_h1.group(1).strip()
                self._advance()
                if first_section:
                    self._render_h1(text)
                continue

            if m_h2:
                text = m_h2.group(1).strip()
                self._advance()
                if first_section:
                    # Skip subtitle H2 (already on cover)
                    if text == self.subtitle:
                        continue
                self._render_h2(text)
                continue

            if m_h3:
                text = m_h3.group(1).strip()
                self._advance()
                self._render_h3(text)
                continue

            if m_h4:
                text = m_h4.group(1).strip()
                self._advance()
                self._render_h4(text)
                continue

            # ── tables ──
            if self._is_table_row(line):
                parsed = self._parse_table()
                if parsed:
                    headers, rows = parsed
                    if not info_table_rendered and self._is_info_table(headers, rows):
                        self._render_info_table(headers, rows)
                        info_table_rendered = True
                    else:
                        self._render_data_table(headers, rows)
                continue

            # ── bullet list ──
            if self._is_bullet(line):
                items = []
                while self.idx < len(self.lines) and self._is_bullet(self.lines[self.idx]):
                    raw = re.sub(r"^[-*]\s+", "", self.lines[self.idx])
                    items.append(raw)
                    self._advance()
                for item in items:
                    p = self.doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    parse_inline(p, item)
                continue

            # ── numbered list ──
            if self._is_numbered(line):
                items = []
                while self.idx < len(self.lines) and self._is_numbered(self.lines[self.idx]):
                    raw = re.sub(r"^\d+\.\s+", "", self.lines[self.idx])
                    items.append(raw)
                    self._advance()
                for item in items:
                    p = self.doc.add_paragraph(style="List Number")
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    parse_inline(p, item)
                continue

            # ── regular paragraph ──
            p = new_paragraph(self.doc)
            parse_inline(p, line)
            self._advance()

        # ── footer ──
        self.doc.add_paragraph()
        p = new_paragraph(self.doc, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0)
        make_run(p, "— Akhir Dokumen —", italic=True, size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))

        # ── save ──
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(self.output_path))
        print(f"✅ DOCX saved to: {self.output_path}")

    @classmethod
    def from_text(cls, text, output_path):
        """Buat DOCX langsung dari string markdown (tanpa file .md di disk)."""
        inst = cls(text=text, output_path=output_path)
        inst.convert()
        return inst.output_path


# ═══════════════════════════════════════════════════════════════════════════════
#  CLI
# ═══════════════════════════════════════════════════════════════════════════════

def find_notulen_md(folder_path):
    """Find the first Notulen_*.md file in a folder."""
    folder = Path(folder_path)
    if not folder.is_dir():
        return None
    candidates = sorted(folder.glob("Notulen*.md"))
    return candidates[0] if candidates else None


def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown notulen to professionally formatted DOCX.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
examples:
  python generate_docx.py --md transcribe_hasil/03_event/Notulen_Event.md
  python generate_docx.py --folder 03_rapat-koordinasi-event
  python generate_docx.py --all
  python generate_docx.py --md file.md --output hasil.docx
        """,
    )
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--md", metavar="PATH",
                       help="Path to a specific Markdown (.md) file")
    group.add_argument("--folder", metavar="NAME",
                       help="Folder name inside transcribe_hasil/ (auto-finds Notulen_*.md)")
    group.add_argument("--all", action="store_true",
                       help="Convert all folders in transcribe_hasil/ that contain Notulen_*.md")
    parser.add_argument("--output", metavar="PATH",
                        help="Custom output .docx path (only with --md)")

    args = parser.parse_args()

    script_dir = Path(__file__).resolve().parent
    hasil_dir = script_dir / "transcribe_hasil"

    # ── Mode: --all ──
    if args.all:
        if not hasil_dir.is_dir():
            print("ERROR: transcribe_hasil/ directory not found.")
            sys.exit(1)

        converted = 0
        for folder in sorted(hasil_dir.iterdir()):
            if not folder.is_dir():
                continue
            md_file = find_notulen_md(folder)
            if md_file:
                print(f"\n📄 Processing: {md_file.name}")
                try:
                    converter = MdToDocx(str(md_file))
                    converter.convert()
                    converted += 1
                except Exception as e:
                    print(f"   ❌ Error: {e}")

        print(f"\n✅ {converted} file(s) converted.")
        return

    # ── Mode: --folder ──
    if args.folder:
        target = hasil_dir / args.folder
        if not target.is_dir():
            print(f"ERROR: folder not found: {target}")
            sys.exit(1)
        md_file = find_notulen_md(target)
        if not md_file:
            print(f"ERROR: no Notulen_*.md file found in {target}")
            sys.exit(1)
        args.md = str(md_file)

    # ── Mode: --md ──
    if not args.md:
        parser.print_help()
        print("\nERROR: specify --md, --folder, or --all")
        sys.exit(1)

    try:
        converter = MdToDocx(args.md, args.output)
        converter.convert()
    except FileNotFoundError as e:
        print(f"ERROR: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"ERROR: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
